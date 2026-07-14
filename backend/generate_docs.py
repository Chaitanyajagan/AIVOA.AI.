import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
import os

def set_cell_background(cell, fill_hex):
    """Sets background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets cell padding (margins) in twentieths of a point (dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_document():
    doc = Document()

    # Set page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Style colors
    primary_color = RGBColor(31, 78, 121)    # Deep Blue
    secondary_color = RGBColor(127, 127, 127) # Grey
    text_color = RGBColor(51, 51, 51)         # Dark Charcoal

    # Title Page
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("AI-First CRM HCP Module\nLog Interaction Screen")
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_run.font.color.rgb = primary_color

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle_p.add_run("Technical & Architecture Documentation\nRound 1 Interview Assignment")
    sub_run.font.name = 'Arial'
    sub_run.font.size = Pt(16)
    sub_run.font.color.rgb = secondary_color

    # Spacer
    for _ in range(3):
        doc.add_paragraph()

    # Candidate Section Box
    candidate_table = doc.add_table(rows=1, cols=1)
    candidate_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = candidate_table.rows[0].cells[0]
    set_cell_background(cell, "F2F2F2")
    set_cell_margins(cell, top=200, bottom=200, left=300, right=300)
    
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("Candidate Profile Details:\n")
    r.font.name = 'Arial'
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.color.rgb = primary_color

    details = [
        ("Candidate Name: ", "Chaitanya Jagan"),
        ("Email Address:  ", "chaitanyajagan.dev@gmail.com"),
        ("LinkedIn:       ", "https://linkedin.com/in/chaitanyajagan"),
        ("GitHub Profile: ", "https://github.com/chaitanyajagan"),
        ("Specialization: ", "AI & Machine Learning Engineering")
    ]
    for label, val in details:
        p2 = cell.add_paragraph()
        run_l = p2.add_run(label)
        run_l.font.name = 'Arial'
        run_l.font.size = Pt(10.5)
        run_l.font.bold = True
        
        run_v = p2.add_run(val)
        run_v.font.name = 'Arial'
        run_v.font.size = Pt(10.5)
        run_v.font.color.rgb = text_color

    doc.add_page_break()

    # Section 1: Executive Summary
    h1 = doc.add_heading("1. Executive Summary & Objective", level=1)
    h1.runs[0].font.color.rgb = primary_color
    h1.runs[0].font.name = 'Arial'
    h1.runs[0].font.bold = True

    p = doc.add_paragraph(
        "This document describes the technical architecture, implementation, and execution details for the "
        "AI-First Customer Relationship Management (CRM) Healthcare Professional (HCP) module. The core "
        "deliverable is the 'Log Interaction Screen', designed to offer life science field representatives the "
        "flexibility to log engagement activities (like meetings, calls, and email discussions) via either "
        "a structured web form or an interactive conversational chat console."
    )
    p.runs[0].font.name = 'Arial'
    p.runs[0].font.size = Pt(11)

    p = doc.add_paragraph(
        "Using LangGraph as the AI Agent Framework and a Groq-powered Large Language Model (LLM), the application "
        "extracts critical entities from the representative's conversations in real-time, pre-filling "
        "the structured input form, recommending contextual follow-up actions, and managing interaction logs."
    )
    p.runs[0].font.name = 'Arial'
    p.runs[0].font.size = Pt(11)

    # Section 2: Tech Stack
    h2 = doc.add_heading("2. Core Requirements & Tech Stack", level=1)
    h2.runs[0].font.color.rgb = primary_color
    h2.runs[0].font.name = 'Arial'
    h2.runs[0].font.bold = True

    stack_table = doc.add_table(rows=7, cols=2)
    stack_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Layer", "Technology Choice & Description"]
    
    # Header Row
    for i, title in enumerate(headers):
        cell = stack_table.rows[0].cells[i]
        set_cell_background(cell, "1F4E79")
        p = cell.paragraphs[0]
        r = p.add_run(title)
        r.font.name = 'Arial'
        r.font.size = Pt(11)
        r.font.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_margins(cell, top=100, bottom=100, left=150, right=150)

    rows = [
        ("Frontend UI", "React (Vite template) with Redux Toolkit for centralized state management."),
        ("Backend Framework", "Python 3.14 with FastAPI to expose RESTful endpoints and chat endpoints."),
        ("AI Agent Framework", "LangGraph for state graphs and structured tool-calling pipelines."),
        ("LLM Integrations", "Groq API client (gemma2-9b-it model) with intelligent rule-based local mock fallback."),
        ("Database Engine", "SQLite for zero-setup local dev; SQLAlchemy ORM for clean relational definitions."),
        ("Fonts & Style", "Google Fonts (Inter) styled using Vanilla CSS with native glassmorphism, transitions and themes.")
    ]

    for idx, (layer, desc) in enumerate(rows):
        row = stack_table.rows[idx + 1]
        
        cell_l = row.cells[0]
        p_l = cell_l.paragraphs[0]
        r_l = p_l.add_run(layer)
        r_l.font.name = 'Arial'
        r_l.font.size = Pt(10)
        r_l.font.bold = True
        set_cell_margins(cell_l, top=80, bottom=80, left=120, right=120)
        
        cell_d = row.cells[1]
        p_d = cell_d.paragraphs[0]
        r_d = p_d.add_run(desc)
        r_d.font.name = 'Arial'
        r_d.font.size = Pt(10)
        set_cell_margins(cell_d, top=80, bottom=80, left=120, right=120)
        if idx % 2 == 1:
            set_cell_background(cell_l, "F9F9F9")
            set_cell_background(cell_d, "F9F9F9")

    doc.add_paragraph()

    # Section 3: Folder Structure
    h3 = doc.add_heading("3. File Directory Structure", level=1)
    h3.runs[0].font.color.rgb = primary_color
    h3.runs[0].font.name = 'Arial'
    h3.runs[0].font.bold = True

    p = doc.add_paragraph("The workspace is divided into two primary monorepo components:")
    p.runs[0].font.name = 'Arial'

    structure_text = (
        "aivoa/\n"
        "├── backend/\n"
        "│   ├── database.py             # SQLite setup, tables (HCP, Material, Sample, Interaction) & SQLAlchemy ORM\n"
        "│   ├── seed.py                 # Seeding script to create initial doctors, studies, and sample inventory\n"
        "│   ├── agent.py                # LangGraph State Graph definitions and the 5 specific DB tools\n"
        "│   ├── main.py                 # FastAPI endpoints, CORS middleware, and Uvicorn runner\n"
        "│   └── requirements.txt        # Python pip package listings\n"
        "├── frontend/\n"
        "│   ├── src/\n"
        "│   │   ├── components/\n"
        "│   │   │   ├── AIAssistant.jsx        # Conversational assistant panel with suggested prompt buttons\n"
        "│   │   │   ├── LogInteractionForm.jsx # Structured log form with simulated audio transcription drawer\n"
        "│   │   │   ├── InteractionHistory.jsx# Feed showing all logged items with inline edits\n"
        "│   │   │   └── SettingsModal.jsx     # Groq API settings modal\n"
        "│   │   ├── store/\n"
        "│   │   │   └── index.js               # Redux Store: slices, actions, and async fetch thunks\n"
        "│   │   ├── App.jsx                    # Core wrapper supporting dark theme state toggling\n"
        "│   │   ├── index.css                  # Style sheet containing standard design variables\n"
        "│   │   └── main.jsx                   # Vite bootstrap entry point\n"
        "│   ├── package.json                   # Web dependencies (Redux, Lucide, standard React)\n"
        "│   └── vite.config.js                 # Vite compile configurations\n"
        "└── README.md                          # Visual architecture and running commands"
    )

    structure_box = doc.add_table(rows=1, cols=1)
    structure_box.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell_s = structure_box.rows[0].cells[0]
    set_cell_background(cell_s, "F4F5F7")
    set_cell_margins(cell_s, top=150, bottom=150, left=200, right=200)
    p_s = cell_s.paragraphs[0]
    r_s = p_s.add_run(structure_text)
    r_s.font.name = 'Courier New'
    r_s.font.size = Pt(9.5)

    doc.add_page_break()

    # Section 4: LangGraph Agent & DB Tools
    h4 = doc.add_heading("4. LangGraph AI Agent & Database Tools", level=1)
    h4.runs[0].font.color.rgb = primary_color
    h4.runs[0].font.name = 'Arial'
    h4.runs[0].font.bold = True

    p = doc.add_paragraph(
        "The core intelligence is driven by a LangGraph State Graph. The graph transitions messages, "
        "tracks an active interaction draft object, and updates suggested actions. It incorporates "
        "the following 5 database-backed tools decorator functions to interact with the SQLite models:"
    )
    p.runs[0].font.name = 'Arial'

    tools_table = doc.add_table(rows=6, cols=3)
    tools_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers_t = ["Tool Name", "Arguments / Schema", "Database Operation & Functionality"]
    for i, title in enumerate(headers_t):
        cell = tools_table.rows[0].cells[i]
        set_cell_background(cell, "1F4E79")
        p = cell.paragraphs[0]
        r = p.add_run(title)
        r.font.name = 'Arial'
        r.font.size = Pt(10.5)
        r.font.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_margins(cell, top=100, bottom=100, left=150, right=150)

    tools = [
        ("search_hcp", "query: str", "Filters HCP records using standard SQL LIKE checks against name, specialty, or hospital locations."),
        ("get_available_materials_and_samples", "None", "Queries available PDFs, brochures, and drug samples alongside active inventory counts."),
        ("get_interaction_history", "hcp_id: int", "Queries the database for all logged interactions linked to the specific HCP ID, ordered by date."),
        ("log_interaction", "hcp_id, interaction_type, date, time, attendees, topics_discussed, materials_shared_ids, samples_distributed_ids, sentiment, outcomes, follow_up_actions", "Creates and saves a new Interaction record in the DB, links materials/samples associations, and decrements sample stocks."),
        ("edit_interaction", "interaction_id: int, updated_fields: dict", "Queries the logged record by ID and applies mutations dynamically to changed attributes in the database.")
    ]

    for idx, (tname, targs, tdesc) in enumerate(tools):
        row = tools_table.rows[idx + 1]
        
        c0 = row.cells[0]
        p0 = c0.paragraphs[0]
        r0 = p0.add_run(tname)
        r0.font.name = 'Arial'
        r0.font.size = Pt(9.5)
        r0.font.bold = True
        set_cell_margins(c0, top=80, bottom=80, left=120, right=120)
        
        c1 = row.cells[1]
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(targs)
        r1.font.name = 'Courier New'
        r1.font.size = Pt(9)
        set_cell_margins(c1, top=80, bottom=80, left=120, right=120)
        
        c2 = row.cells[2]
        p2 = c2.paragraphs[0]
        r2 = p2.add_run(tdesc)
        r2.font.name = 'Arial'
        r2.font.size = Pt(9.5)
        set_cell_margins(c2, top=80, bottom=80, left=120, right=120)
        
        if idx % 2 == 1:
            for c in [c0, c1, c2]:
                set_cell_background(c, "F9F9F9")

    doc.add_paragraph()

    # Section 5: Real-time Form Sync
    h5 = doc.add_heading("5. Real-time Form Synchronization Flow", level=1)
    h5.runs[0].font.color.rgb = primary_color
    h5.runs[0].font.name = 'Arial'
    h5.runs[0].font.bold = True

    p = doc.add_paragraph(
        "One of the standout design elements is the real-time synchronization between the structured form "
        "and the AI Assistant conversation. When a representative chats about their day (e.g. 'I met Dr. Smith, "
        "discussed CardioLife, positive sentiment, left some 5mg samples'):\n"
        "1. The FastAPI chat route sends the input array and current form draft to the LangGraph node.\n"
        "2. The AI agent analyzes the text and invokes search tools to resolve named entities to IDs.\n"
        "3. The agent calls update_draft_interaction(fields) to output the extracted fields.\n"
        "4. The backend server packs these fields and sends them in the JSON payload response.\n"
        "5. The Redux frontend store interceptor merges the extracted fields back into the local state, "
        "causing the React structured form fields to update instantly."
    )
    p.runs[0].font.name = 'Arial'
    p.runs[0].font.size = Pt(11)

    # Section 6: Setup Guide
    h6 = doc.add_heading("6. Setup & Execution Instructions", level=1)
    h6.runs[0].font.color.rgb = primary_color
    h6.runs[0].font.name = 'Arial'
    h6.runs[0].font.bold = True

    doc.add_heading("Backend Setup", level=2)
    p = doc.add_paragraph(
        "1. Change directory to backend folder: cd backend\n"
        "2. Create virtual environment: python -m venv venv\n"
        "3. Activate environment: .\\venv\\Scripts\\activate\n"
        "4. Install dependencies: pip install -r requirements.txt\n"
        "5. Seed the SQLite database: python seed.py\n"
        "6. Start the API server: python main.py"
    )
    p.runs[0].font.name = 'Arial'

    doc.add_heading("Frontend Setup", level=2)
    p = doc.add_paragraph(
        "1. Change directory to frontend folder: cd frontend\n"
        "2. Install packages: npm install\n"
        "3. Start dev server: npm run dev"
    )
    p.runs[0].font.name = 'Arial'

    # Save document
    output_path = os.path.abspath(os.path.join("..", "AI-First_CRM_HCP_Module_Documentation.docx"))
    doc.save(output_path)
    print(f"Document saved successfully at: {output_path}")

if __name__ == "__main__":
    create_document()
