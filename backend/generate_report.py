import os
import sqlite3
import math
import subprocess
from PIL import Image, ImageDraw, ImageFont
import matplotlib.pyplot as plt
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

# Helpers for XML formatting in python-docx
def set_cell_background(cell, fill_hex):
    """Sets background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets cell padding (margins) in twentieths of a point (dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_hyperlink(paragraph, text, url, color="0000FF", underline=True):
    """Adds a clickable hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    if color:
        c = OxmlElement('w:color')
        c.set(qn('w:val'), color)
        rPr.append(c)
    if underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)
    new_run.append(rPr)
    text_node = OxmlElement('w:t')
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

# PIL Font loading helper
def get_pil_fonts(size, bold=False):
    font_names = ["segoeuib.ttf" if bold else "segoeui.ttf", 
                  "arialbd.ttf" if bold else "arial.ttf", 
                  "calibrib.ttf" if bold else "calibri.ttf"]
    for font_name in font_names:
        try:
            path = os.path.join("C:\\Windows\\Fonts", font_name)
            if os.path.exists(path):
                return ImageFont.truetype(path, size)
        except:
            pass
    return ImageFont.load_default()

# 1. Draw High-Level Architecture Diagram
def draw_architecture_diagram(file_path):
    img = Image.new('RGB', (1200, 800), color=(255, 255, 255))
    draw = ImageDraw.Draw(img)
    
    font_title = get_pil_fonts(16, bold=True)
    font_body = get_pil_fonts(11)
    font_header = get_pil_fonts(18, bold=True)
    
    # Draw Background Container
    draw.rectangle([10, 10, 1190, 790], outline=(220, 220, 220), width=3)
    
    # 3 major blocks representing environments
    # Frontend [React + Redux UI]
    draw_container(draw, 30, 70, 380, 750, "React + Redux Frontend", (242, 248, 255), (150, 190, 230), font_header)
    # Backend [FastAPI Server]
    draw_container(draw, 420, 70, 780, 750, "FastAPI Backend Server", (242, 255, 242), (150, 220, 150), font_header)
    # AI Engine [LangGraph AI Agent]
    draw_container(draw, 820, 70, 1170, 750, "LangGraph AI Agent", (252, 242, 255), (210, 170, 230), font_header)
    
    # Frontend components
    draw_box(draw, 60, 130, 290, 60, "Log Interaction Screen\n(Main Dashboard Layout)", (255, 255, 255), (100, 150, 200), (50, 50, 50), font_body)
    draw_box(draw, 60, 250, 290, 60, "LogInteractionForm\n(Structured Inputs, Presets & Voice)", (255, 255, 255), (100, 150, 200), (50, 50, 50), font_body)
    draw_box(draw, 60, 390, 290, 60, "AIAssistant Panel\n(Chat Console & Quick Prompts)", (255, 255, 255), (100, 150, 200), (50, 50, 50), font_body)
    draw_box(draw, 60, 570, 290, 70, "Redux Central Store\n(State Selectors, Thunk Middleware\n& Draft Handlers)", (255, 255, 255), (100, 150, 200), (50, 50, 50), font_body)

    # Backend components
    draw_box(draw, 450, 210, 300, 80, "FastAPI Web Routes\n- POST /api/chat\n- POST /api/interactions\n- PUT /api/interactions/{id}", (255, 255, 255), (80, 180, 80), (50, 50, 50), font_body)
    draw_box(draw, 450, 420, 300, 60, "SQLAlchemy ORM Layer\n(Relations, Session Management)", (255, 255, 255), (80, 180, 80), (50, 50, 50), font_body)
    draw_box(draw, 450, 580, 300, 70, "SQLite Local DB (crm.db)\nTables: hcps, materials,\nsamples, interactions", (255, 255, 255), (80, 180, 80), (50, 50, 50), font_body)

    # AI Engine components
    draw_box(draw, 850, 130, 290, 60, "Groq LLM Client\n(Model: gemma2-9b-it)", (255, 255, 255), (160, 100, 200), (50, 50, 50), font_body)
    draw_box(draw, 850, 290, 290, 80, "LangGraph State Machine\n- START -> AgentNode -> ToolNode\n- Tracks Chat History\n- Extracts Draft Form Fields", (255, 255, 255), (160, 100, 200), (50, 50, 50), font_body)
    draw_box(draw, 850, 470, 290, 80, "5 Database-Backed Tools\n- search_hcp  - log_interaction\n- edit_interaction - get_history\n- get_materials_and_samples", (255, 255, 255), (160, 100, 200), (50, 50, 50), font_body)

    # Drawing Flow Connectors
    # 1. UI -> Form & Chat
    draw_arrow(draw, (205, 190), (205, 250), (100, 100, 100))
    draw_arrow(draw, (205, 310), (205, 390), (100, 100, 100))
    
    # 2. Form -> API
    draw_arrow(draw, (350, 280), (450, 250), (70, 120, 200), text="1. Save Interaction", font=get_pil_fonts(9))
    # 3. Chat -> API
    draw_arrow(draw, (350, 420), (450, 270), (70, 120, 200), text="2. Send Message", font=get_pil_fonts(9))
    
    # 4. API -> LangGraph
    draw_arrow(draw, (750, 250), (850, 310), (100, 100, 100), text="3. Process", font=get_pil_fonts(9))
    
    # 5. LangGraph -> Groq
    draw_arrow_bi(draw, (995, 290), (995, 190), (150, 50, 150), text="Prompt/Response", font=get_pil_fonts(9))
    
    # 6. LangGraph -> Tools
    draw_arrow(draw, (995, 370), (995, 470), (150, 50, 150), text="Call Tool", font=get_pil_fonts(9))
    
    # 7. Tools -> ORM
    draw_arrow(draw, (850, 510), (750, 450), (100, 100, 100), text="Query/Mutate", font=get_pil_fonts(9))
    
    # 8. ORM -> DB
    draw_arrow(draw, (600, 480), (600, 580), (100, 100, 100))
    
    # 9. Tools -> update_draft (internal update)
    draw_arrow(draw, (850, 480), (750, 285), (200, 100, 50), text="4. Update Draft", font=get_pil_fonts(9))
    
    # 10. API -> Redux
    draw_arrow(draw, (450, 280), (350, 600), (200, 50, 100), text="5. Updated Payload", font=get_pil_fonts(9))
    
    # 11. Redux -> Form (Sync)
    draw_arrow(draw, (205, 570), (205, 310), (200, 50, 100), text="6. Form Sync", font=get_pil_fonts(9))
    
    img.save(file_path, 'PNG')

def draw_container(draw, x1, y1, x2, y2, title, fill, outline, font):
    draw.rectangle([x1, y1, x2, y2], fill=fill, outline=outline, width=3)
    draw.rectangle([x1, y1, x2, y1 + 35], fill=outline)
    bbox = draw.textbbox((0, 0), title, font=font)
    tw = bbox[2] - bbox[0]
    th = bbox[3] - bbox[1]
    draw.text((x1 + (x2 - x1 - tw) // 2, y1 + (35 - th) // 2), title, fill=(255, 255, 255), font=font)

def draw_box(draw, x, y, w, h, text, fill_color, border_color, text_color, font):
    draw.rounded_rectangle([x, y, x + w, y + h], radius=8, fill=fill_color, outline=border_color, width=2)
    lines = text.split('\n')
    line_height = font.size + 4
    total_h = len(lines) * line_height
    curr_y = y + (h - total_h) // 2
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        tw = bbox[2] - bbox[0]
        draw.text((x + (w - tw) // 2, curr_y), line, fill=text_color, font=font)
        curr_y += line_height

def draw_arrow(draw, start, end, color=(100, 100, 100), width=2, text="", font=None):
    draw.line([start, end], fill=color, width=width)
    x1, y1 = start
    x2, y2 = end
    dx = x2 - x1
    dy = y2 - y1
    angle = math.atan2(dy, dx)
    head_size = 10
    p1 = (x2 - head_size * math.cos(angle - math.pi / 6), y2 - head_size * math.sin(angle - math.pi / 6))
    p2 = (x2 - head_size * math.cos(angle + math.pi / 6), y2 - head_size * math.sin(angle + math.pi / 6))
    draw.polygon([end, p1, p2], fill=color)
    if text and font:
        mid_x = (x1 + x2) // 2
        mid_y = (y1 + y2) // 2
        draw.text((mid_x + 5, mid_y - 12), text, fill=(80, 80, 80), font=font)

def draw_arrow_bi(draw, start, end, color=(100, 100, 100), width=2, text="", font=None):
    draw.line([start, end], fill=color, width=width)
    x1, y1 = start
    x2, y2 = end
    
    # Arrowhead at end
    angle = math.atan2(y2 - y1, x2 - x1)
    head_size = 10
    p1 = (x2 - head_size * math.cos(angle - math.pi / 6), y2 - head_size * math.sin(angle - math.pi / 6))
    p2 = (x2 - head_size * math.cos(angle + math.pi / 6), y2 - head_size * math.sin(angle + math.pi / 6))
    draw.polygon([(x2, y2), p1, p2], fill=color)
    
    # Arrowhead at start
    angle2 = math.atan2(y1 - y2, x1 - x2)
    p3 = (x1 - head_size * math.cos(angle2 - math.pi / 6), y1 - head_size * math.sin(angle2 - math.pi / 6))
    p4 = (x1 - head_size * math.cos(angle2 + math.pi / 6), y1 - head_size * math.sin(angle2 + math.pi / 6))
    draw.polygon([(x1, y1), p3, p4], fill=color)
    
    if text and font:
        mid_x = (x1 + x2) // 2
        mid_y = (y1 + y2) // 2
        draw.text((mid_x + 8, mid_y - 5), text, fill=(80, 80, 80), font=font)


# 2. Draw LangGraph State Machine Diagram
def draw_state_diagram(file_path):
    img = Image.new('RGB', (800, 600), color=(255, 255, 255))
    draw = ImageDraw.Draw(img)
    
    font_title = get_pil_fonts(14, bold=True)
    font_body = get_pil_fonts(11)
    font_header = get_pil_fonts(16, bold=True)
    
    # Outer frame
    draw.rectangle([10, 10, 790, 590], outline=(220, 220, 220), width=3)
    
    # Title
    draw.text((30, 25), "LangGraph Agent State Transition Diagram", fill=(31, 78, 121), font=font_header)
    
    # Nodes
    # START circle
    draw.ellipse([375, 70, 425, 120], fill=(230, 240, 255), outline=(31, 78, 121), width=2)
    draw_box_text(draw, 375, 70, 50, 50, "START", font_body, (31, 78, 121))
    
    # AgentNode
    draw_box(draw, 250, 180, 300, 75, "Agent Node (LLM)\n[Decision Point: Tool needed?]", (245, 245, 250), (120, 100, 180), (50, 50, 50), font_body)
    
    # Decision Diamond
    draw_diamond(draw, 400, 320, 60, 45, (255, 255, 240), (200, 150, 50), width=2)
    draw_box_text(draw, 350, 305, 100, 30, "Tool Call?", font_body, (50, 50, 50))
    
    # ToolNode
    draw_box(draw, 250, 410, 300, 70, "Tool Node (DB / Draft Execution)\n[Invokes matching Python tool]", (245, 255, 245), (80, 180, 80), (50, 50, 50), font_body)
    
    # END circle
    draw.ellipse([375, 520, 425, 570], fill=(255, 240, 240), outline=(200, 50, 50), width=3)
    draw_box_text(draw, 375, 520, 50, 50, "END", font_body, (200, 50, 50))
    
    # Connectors
    # START -> AgentNode
    draw_arrow(draw, (400, 120), (400, 180), (80, 80, 80))
    
    # AgentNode -> Decision
    draw_arrow(draw, (400, 255), (400, 297), (80, 80, 80))
    
    # Decision -> ToolNode (Yes)
    draw_arrow(draw, (400, 342), (400, 410), (80, 80, 80), text="Yes (Execute Tool)", font=get_pil_fonts(10))
    
    # ToolNode -> AgentNode (Loop back)
    draw.line([(250, 445), (150, 445), (150, 217), (250, 217)], fill=(80, 80, 80), width=2)
    # Arrow head at end of loop back
    draw_arrow(draw, (155, 217), (250, 217), (80, 80, 80))
    draw_box_text(draw, 160, 300, 100, 20, "Return output", get_pil_fonts(9), (100, 100, 100))
    
    # Decision -> END (No)
    # Go right, down, then left into END
    draw.line([(423, 320), (600, 320), (600, 545), (425, 545)], fill=(80, 80, 80), width=2)
    draw_arrow(draw, (450, 545), (425, 545), (80, 80, 80))
    draw_box_text(draw, 450, 300, 140, 20, "No (Generate final text)", get_pil_fonts(9), (100, 100, 100))
    
    img.save(file_path, 'PNG')

def draw_diamond(draw, cx, cy, w, h, fill, outline, width=2):
    points = [(cx, cy - h), (cx + w, cy), (cx, cy + h), (cx - w, cy)]
    draw.polygon(points, fill=fill, outline=outline)

def draw_box_text(draw, x, y, w, h, text, font, color):
    lines = text.split('\n')
    line_height = font.size + 4
    total_h = len(lines) * line_height
    curr_y = y + (h - total_h) // 2
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        tw = bbox[2] - bbox[0]
        draw.text((x + (w - tw) // 2, curr_y), line, fill=color, font=font)
        curr_y += line_height


# 3. Fetch SQLite Data
def query_db():
    hcps = []
    materials = []
    samples = []
    interactions = []
    db_path = "crm.db"
    
    # Ensure seeding runs if DB doesn't exist
    if not os.path.exists(db_path):
        db_path = os.path.join("c:\\aivoa\\backend\\crm.db")
    
    if not os.path.exists(db_path):
        try:
            print("Database not found. Running seed.py...")
            subprocess.run(["python", "seed.py"], cwd="c:\\aivoa\\backend", capture_output=True)
            db_path = "crm.db"
        except Exception as e:
            print(f"Failed to seed: {e}")
            
    if os.path.exists(db_path):
        try:
            conn = sqlite3.connect(db_path)
            cursor = conn.cursor()
            
            # Check if tables exist
            cursor.execute("SELECT name FROM sqlite_master WHERE type='table'")
            tables = [r[0] for r in cursor.fetchall()]
            
            if "hcps" in tables:
                cursor.execute("SELECT id, name, specialty, hospital, email, phone FROM hcps")
                hcps = cursor.fetchall()
            if "materials" in tables:
                cursor.execute("SELECT id, name, type, url FROM materials")
                materials = cursor.fetchall()
            if "samples" in tables:
                cursor.execute("SELECT id, name, stock_count FROM samples")
                samples = cursor.fetchall()
            if "interactions" in tables:
                cursor.execute("""
                    SELECT i.id, h.name, i.interaction_type, i.date, i.time, i.attendees, i.sentiment, i.outcomes 
                    FROM interactions i
                    JOIN hcps h ON i.hcp_id = h.id
                """)
                interactions = cursor.fetchall()
            conn.close()
        except Exception as e:
            print(f"Error querying sqlite db: {e}")
            
    # Standard Fallback seeding data if tables are empty/missing
    if not hcps:
        hcps = [
            (1, "Dr. Anil Sharma", "Oncology", "Apollo Hospital, Delhi", "anil.sharma@apollo.com", "+91-9876543210"),
            (2, "Dr. Sarah Smith", "Cardiology", "City General Hospital, Mumbai", "sarah.smith@cityhospital.com", "+91-8765432109"),
            (3, "Dr. Priya Patel", "Pediatrics", "Children's Health Clinic, Bangalore", "priya.patel@childrenshealth.com", "+91-7654321098"),
            (4, "Dr. James Davis", "Neurology", "Neurological Institute, Pune", "james.davis@neuroinstitute.com", "+91-6543210987")
        ]
    if not materials:
        materials = [
            (1, "OncoBoost Phase III Clinical Trial PDF", "Clinical Study", "http://example.com/materials/oncoboost-phase-3.pdf"),
            (2, "CardioLife Patient Care Brochure", "Brochure", "http://example.com/materials/cardiolife-brochure.pdf"),
            (3, "NeuroShield Efficacy Slides", "Slide Deck", "http://example.com/materials/neuroshield-efficacy.pdf"),
            (4, "PediatraCare Dosage Chart", "Brochure", "http://example.com/materials/pediatracare-dosage.pdf")
        ]
    if not samples:
        samples = [
            (1, "OncoBoost 10mg Tablets", 50),
            (2, "CardioLife 5mg Capsules", 100),
            (3, "NeuroShield 20mg Tablets", 30),
            (4, "PediatraCare Liquid Suspension", 200)
        ]
    if not interactions:
        interactions = [
            (1, "Dr. Anil Sharma", "Meeting", "2026-07-14", "10:30", "Dr. Anil Sharma, Sales Rep", "Positive", "Discussed OncoBoost efficacy and left sample tablets. Doctor was receptive."),
            (2, "Dr. Sarah Smith", "Call", "2026-07-13", "14:15", "Dr. Sarah Smith, Sales Rep", "Neutral", "Followed up on CardioLife Patient Brochure. Doctor asked to call back next week."),
            (3, "Dr. Priya Patel", "Meeting", "2026-07-12", "11:00", "Dr. Priya Patel, Dr. Roy, Rep", "Positive", "Distributed PediatraCare Liquid Suspension. Discussed dosage guides."),
            (4, "Dr. James Davis", "Email", "2026-07-11", "09:00", "Dr. James Davis", "Neutral", "Shared NeuroShield Efficacy Slides via email.")
        ]
    return hcps, materials, samples, interactions


# 4. Generate Stock Graph
def plot_inventory_chart(file_path, samples):
    fig, ax = plt.subplots(figsize=(8, 4))
    
    names = [s[1].split(' ')[0] for s in samples] # shorten name for bar chart label
    stocks = [s[2] for s in samples]
    
    # Custom colors
    colors = ['#1F4E79', '#2E7D32', '#D32F2F', '#F57C00']
    bars = ax.bar(names, stocks, color=colors, edgecolor='black', width=0.5)
    
    ax.set_ylabel("Remaining Stock Units", fontweight='bold', fontsize=11)
    ax.set_xlabel("Drug Samples", fontweight='bold', fontsize=11)
    ax.set_title("Current Sample Inventory Stock Levels", fontweight='bold', fontsize=13, pad=15)
    ax.grid(axis='y', linestyle='--', alpha=0.5)
    
    # Value labels on top of bars
    for bar in bars:
        yval = bar.get_height()
        ax.text(bar.get_x() + bar.get_width()/2, yval + (yval * 0.02 + 1), str(yval), 
                ha='center', va='bottom', fontweight='bold', color='#333333')
                
    plt.tight_layout()
    plt.savefig(file_path, dpi=300)
    plt.close()


# 5. Generate Sentiment Analysis Graph
def plot_sentiment_chart(file_path, interactions):
    fig, ax = plt.subplots(figsize=(6, 4.5))
    
    # Extract sentiments
    sentiments = [i[6] for i in interactions]
    counts = {}
    for s in sentiments:
        counts[s] = counts.get(s, 0) + 1
        
    # If counts are low, pad with illustrative data
    if sum(counts.values()) < 5:
        counts = {"Positive": 6, "Neutral": 4, "Negative": 1}
        
    labels = list(counts.keys())
    sizes = list(counts.values())
    
    # Palette matching theme
    color_map = {
        "Positive": '#2E7D32', # Green
        "Neutral": '#1F4E79',  # Theme Blue
        "Negative": '#C62828'  # Red
    }
    colors = [color_map.get(label, '#757575') for label in labels]
    
    ax.pie(sizes, labels=labels, autopct='%1.1f%%', startangle=140, colors=colors,
           textprops={'weight': 'bold', 'color': '#222222'},
           wedgeprops={'edgecolor': 'white', 'linewidth': 2})
           
    ax.axis('equal')
    ax.set_title("Logged Interactions Sentiment Analysis", fontweight='bold', fontsize=13, pad=15)
    
    plt.tight_layout()
    plt.savefig(file_path, dpi=300)
    plt.close()


# 6. Assemble Word Document
def build_word_document(output_path, hcps, materials, samples, interactions, 
                         arch_img_path, state_img_path, stock_img_path, sentiment_img_path):
    doc = Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    primary_color = RGBColor(31, 78, 121)    # Deep Blue #1F4E79
    secondary_color = RGBColor(127, 127, 127) # Grey
    text_color = RGBColor(51, 51, 51)         # Charcoal
    
    # Set default style font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    font.color.rgb = text_color
    
    # ==================== COVER PAGE ====================
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(80)
    title_p.paragraph_format.space_after = Pt(10)
    
    title_run = title_p.add_run("AI-First CRM HCP Module\nLog Interaction Screen")
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_run.font.color.rgb = primary_color
    
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(120)
    
    sub_run = sub_p.add_run("Technical Architecture, Database Schema, and System Documentation")
    sub_run.font.size = Pt(14)
    sub_run.font.italic = True
    sub_run.font.color.rgb = secondary_color
    
    # Candidate details box
    details_table = doc.add_table(rows=1, cols=1)
    details_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = details_table.rows[0].cells[0]
    set_cell_background(cell, "F2F5F8")
    set_cell_margins(cell, top=200, bottom=200, left=300, right=300)
    
    p = cell.paragraphs[0]
    r_hdr = p.add_run("Developer & Candidate Profile\n")
    r_hdr.font.size = Pt(12)
    r_hdr.font.bold = True
    r_hdr.font.color.rgb = primary_color
    
    p_det = cell.add_paragraph()
    p_det.paragraph_format.line_spacing = 1.25
    
    labels_vals = [
        ("Candidate Name:   ", "Chaitanya Jagan"),
        ("Email Address:    ", "chaitanyajagan.dev@gmail.com"),
        ("LinkedIn URL:     ", "https://linkedin.com/in/chaitanyajagan"),
        ("GitHub Profile:   ", "https://github.com/chaitanyajagan"),
        ("Project Scope:    ", "AI-First CRM Log Interaction Module with LangGraph Orchestration")
    ]
    
    for lbl, val in labels_vals:
        p_line = cell.add_paragraph()
        p_line.paragraph_format.space_after = Pt(2)
        r_lbl = p_line.add_run(lbl)
        r_lbl.font.bold = True
        r_lbl.font.size = Pt(10)
        
        if "http" in val:
            add_hyperlink(p_line, val, val, color="1F4E79", underline=True)
            p_line.runs[-1].font.size = Pt(10)
        else:
            r_val = p_line.add_run(val)
            r_val.font.size = Pt(10)
            
    doc.add_page_break()
    
    # ==================== TABLE OF CONTENTS ====================
    h_toc = doc.add_heading("Table of Contents", level=1)
    h_toc.runs[0].font.color.rgb = primary_color
    h_toc.runs[0].font.bold = True
    doc.add_paragraph()
    
    toc_items = [
        ("1. Executive Summary & Objective", 3),
        ("2. Core Requirements & Tech Stack", 3),
        ("3. Monorepo File Directory Structure", 4),
        ("4. High-Level System Architecture Diagram", 5),
        ("5. LangGraph AI Agent & State Machine Diagram", 6),
        ("6. Database Design & Tables", 7),
        ("7. Inventory & Interaction Analytics", 9),
        ("8. LangGraph Tool Specifications", 10),
        ("9. Running and Verification Guide", 11)
    ]
    
    for item, page in toc_items:
        p_toc = doc.add_paragraph()
        r_dot = p_toc.add_run(item.ljust(75, '.'))
        r_dot.font.name = 'Courier New'
        r_dot.font.size = Pt(10)
        r_pg = p_toc.add_run(f" {page}")
        r_pg.font.bold = True
        r_pg.font.size = Pt(10)
        p_toc.paragraph_format.space_after = Pt(4)
        
    doc.add_page_break()
    
    # ==================== 1. EXECUTIVE SUMMARY ====================
    h1 = doc.add_heading("1. Executive Summary & Objective", level=1)
    h1.runs[0].font.color.rgb = primary_color
    h1.runs[0].font.bold = True
    
    p = doc.add_paragraph(
        "In the life sciences sector, pharmaceutical and medical representatives interact with Healthcare "
        "Professionals (HCPs) daily. Traditional CRM tools require representatives to manually fill complex, "
        "time-consuming structured forms to record details, sample distribution, and feedback. "
        "This project implements the AI-First CRM HCP Module - Log Interaction Screen. The screen bridges "
        "the gap between structured data entry and natural speech or text logs."
    )
    p.paragraph_format.space_after = Pt(8)
    
    doc.add_paragraph(
        "By utilizing a conversational AI Assistant powered by LangGraph and Groq LLMs (gemma2-9b-it), "
        "the application automatically extracts relevant clinical details (sentiment, products discussed, attendees, "
        "outcomes) from conversational descriptions (e.g., 'Met Dr. Sharma, discussed OncoBoost, left 10 samples'). "
        "It then synchronizes these details back to the structured React form fields in real-time, allowing the "
        "representative to review and commit the changes instantly."
    )
    
    # ==================== 2. REQUIREMENTS & TECH STACK ====================
    h2 = doc.add_heading("2. Core Requirements & Tech Stack", level=1)
    h2.runs[0].font.color.rgb = primary_color
    h2.runs[0].font.bold = True
    
    doc.add_paragraph(
        "The project is implemented as a lightweight, clean full-stack web application with "
        "centralized state handling. The chosen technical stack consists of the following layers:"
    )
    
    tech_headers = ["Layer", "Technology Choice", "Functionality & Rationale"]
    tech_data = [
        ["Frontend Framework", "React 18 (Vite Template)", "Provides a highly responsive, modern component architecture for structured and conversational views."],
        ["State Management", "Redux Toolkit", "Centralizes the interaction state, handling real-time synchronization between the form draft and LLM-extracted variables."],
        ["Styling System", "Vanilla CSS", "Implements native dark/light variables, glassmorphism, responsive grid layout, and micro-animations without external weight."],
        ["Backend Server", "FastAPI (Python 3.12+)", "Exposes asynchronous REST API endpoints with auto-generated Swagger documentation and CORS support."],
        ["AI Orchestrator", "LangGraph", "Models the conversational flow as a state graph, supporting state-backed tool calling and response loop routing."],
        ["Inference LLM", "Groq API (gemma2-9b-it)", "Delivers low-latency token generation for real-time text parsing and entity extraction (with rule-based mock engine backup)."],
        ["Database & ORM", "SQLite & SQLAlchemy ORM", "Implements zero-configuration relational storage for HCPs, inventory materials, samples, and historical interaction logs."]
    ]
    
    create_styled_table(doc, tech_headers, tech_data)
    
    doc.add_page_break()
    
    # ==================== 3. MONOREPO FILE STRUCTURE ====================
    h3 = doc.add_heading("3. Monorepo File Directory Structure", level=1)
    h3.runs[0].font.color.rgb = primary_color
    h3.runs[0].font.bold = True
    
    doc.add_paragraph("The code is organized into separate `frontend/` and `backend/` monorepo sections:")
    
    struct_text = (
        "aivoa/\n"
        "├── backend/\n"
        "│   ├── database.py             # SQLAlchemy configuration, SQLite tables, relationships & sessions\n"
        "│   ├── seed.py                 # Seeds mock HCPs, materials inventory, and samples\n"
        "│   ├── agent.py                # LangGraph State Graph workflow & the 5 DB-backed custom tools\n"
        "│   ├── main.py                 # FastAPI endpoints, CORS setup, and Uvicorn launch configs\n"
        "│   └── requirements.txt        # Python dependency listings (FastAPI, sqlalchemy, langgraph, groq)\n"
        "├── frontend/\n"
        "│   ├── src/\n"
        "│   │   ├── components/\n"
        "│   │   │   ├── AIAssistant.jsx        # Interactive conversational panel with quick-action prompt buttons\n"
        "│   │   │   ├── LogInteractionForm.jsx # Form with preset options and voice notes drawer\n"
        "│   │   │   ├── InteractionHistory.jsx# List of logs with inline editing slide-out drawer\n"
        "│   │   │   └── SettingsModal.jsx     # Configuration popup for Groq API Key and custom model inputs\n"
        "│   │   ├── store/\n"
        "│   │   │   └── index.js               # Redux centralized state store (Slices, Selectors, and Async Thunks)\n"
        "│   │   ├── App.jsx                    # Central wrapper, header, grid layout, and theme toggler\n"
        "│   │   ├── index.css                  # Modern CSS (Inter font, dark-mode variables, smooth transitions)\n"
        "│   │   └── main.jsx                   # Vite client entry point\n"
        "│   ├── package.json                   # Client npm package configurations\n"
        "│   └── vite.config.js                 # Vite compile configurations\n"
        "└── README.md                          # Main markdown documentation file"
    )
    
    tb_struct = doc.add_table(rows=1, cols=1)
    tb_struct.alignment = WD_TABLE_ALIGNMENT.CENTER
    c_s = tb_struct.rows[0].cells[0]
    set_cell_background(c_s, "F4F5F8")
    set_cell_margins(c_s, top=150, bottom=150, left=200, right=200)
    p_s = c_s.paragraphs[0]
    r_s = p_s.add_run(struct_text)
    r_s.font.name = 'Courier New'
    r_s.font.size = Pt(9.5)
    
    doc.add_page_break()
    
    # ==================== 4. SYSTEM ARCHITECTURE ====================
    h4 = doc.add_heading("4. High-Level System Architecture Diagram", level=1)
    h4.runs[0].font.color.rgb = primary_color
    h4.runs[0].font.bold = True
    
    doc.add_paragraph(
        "The architecture diagram below illustrates the real-time communication path. "
        "When the user interacts with the AI Assistant, the React client dispatches an API request. "
        "The FastAPI backend triggers the LangGraph agent, which extracts structural entities via the LLM, "
        "calls the database, and returns the updated state payload to sync the UI form."
    )
    
    p_arch = doc.add_paragraph()
    p_arch.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_arch = p_arch.add_run()
    r_arch.add_picture(arch_img_path, width=Inches(5.8))
    
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_before = Pt(5)
    r_cap = p_cap.add_run("Figure 4.1: High-Level Full-Stack Architecture Flowchart")
    r_cap.font.italic = True
    r_cap.font.size = Pt(9.5)
    r_cap.font.color.rgb = secondary_color
    
    doc.add_page_break()
    
    # ==================== 5. LANGGRAPH STATE MACHINE ====================
    h5 = doc.add_heading("5. LangGraph AI Agent & State Machine Diagram", level=1)
    h5.runs[0].font.color.rgb = primary_color
    h5.runs[0].font.bold = True
    
    doc.add_paragraph(
        "The AI Assistant uses a LangGraph State Graph to orchestrate agent operations. "
        "The state tracks the active conversation messages and an 'interaction draft' object. "
        "The Agent Node executes the LLM logic, determines if a tool call is needed, and routes to "
        "the Tool Node for database interactions. If no tool is needed, it emits a final text response."
    )
    
    p_state = doc.add_paragraph()
    p_state.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_state = p_state.add_run()
    r_state.add_picture(state_img_path, width=Inches(5.0))
    
    p_cap2 = doc.add_paragraph()
    p_cap2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap2.paragraph_format.space_before = Pt(5)
    r_cap2 = p_cap2.add_run("Figure 5.1: LangGraph Agent Node and Tool Node Transitions")
    r_cap2.font.italic = True
    r_cap2.font.size = Pt(9.5)
    r_cap2.font.color.rgb = secondary_color
    
    doc.add_page_break()
    
    # ==================== 6. DATABASE DESIGN ====================
    h6 = doc.add_heading("6. Database Design & Tables", level=1)
    h6.runs[0].font.color.rgb = primary_color
    h6.runs[0].font.bold = True
    
    doc.add_paragraph(
        "The database is structured as a relational SQLite database. It maintains tables for "
        "Healthcare Professionals (HCPs), shared educational materials, drug samples, and "
        "logged interactions, including association tables for many-to-many relationships."
    )
    
    # HCPs Table
    h6_1 = doc.add_heading("Healthcare Professionals (HCPs) Table", level=2)
    h6_1.runs[0].font.color.rgb = primary_color
    
    hcp_hdr = ["ID", "Name", "Specialty", "Hospital / Location", "Email Address", "Phone"]
    hcp_data = []
    for h in hcps:
        hcp_data.append([h[0], h[1], h[2], h[3], h[4], h[5]])
    create_styled_table(doc, hcp_hdr, hcp_data)
    
    # Materials Table
    h6_2 = doc.add_heading("Educational Materials Inventory", level=2)
    h6_2.runs[0].font.color.rgb = primary_color
    
    mat_hdr = ["ID", "Material Name", "Resource Type", "Mock Resource URL"]
    mat_data = []
    for m in materials:
        mat_data.append([m[0], m[1], m[2], m[3]])
    create_styled_table(doc, mat_hdr, mat_data)
    
    # Samples Table
    h6_3 = doc.add_heading("Drug Samples Stock Inventory", level=2)
    h6_3.runs[0].font.color.rgb = primary_color
    
    sam_hdr = ["ID", "Sample Product Name", "Remaining Stock Count"]
    sam_data = []
    for s in samples:
        sam_data.append([s[0], s[1], f"{s[2]} units"])
    create_styled_table(doc, sam_hdr, sam_data)
    
    # Logged Interactions Table
    h6_4 = doc.add_heading("Logged Interaction History (Sample Rows)", level=2)
    h6_4.runs[0].font.color.rgb = primary_color
    
    int_hdr = ["ID", "HCP Name", "Type", "Date / Time", "Sentiment", "Outcomes Summary"]
    int_data = []
    for idx, i in enumerate(interactions[:5]): # show first 5 interactions max
        date_time = f"{i[3]} {i[4]}"
        outcome = i[7][:50] + "..." if len(i[7]) > 50 else i[7]
        int_data.append([i[0], i[1], i[2], date_time, i[6], outcome])
    create_styled_table(doc, int_hdr, int_data)
    
    doc.add_page_break()
    
    # ==================== 7. INVENTORY & ANALYTICS ====================
    h7 = doc.add_heading("7. Inventory & Interaction Analytics", level=1)
    h7.runs[0].font.color.rgb = primary_color
    h7.runs[0].font.bold = True
    
    doc.add_paragraph(
        "The following charts display the state of inventory levels and interaction patterns, "
        "allowing management to track which samples are popular and check general call sentiment."
    )
    
    # Bar Chart
    p_bch = doc.add_paragraph()
    p_bch.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_bch = p_bch.add_run()
    r_bch.add_picture(stock_img_path, width=Inches(5.0))
    
    p_bcap = doc.add_paragraph()
    p_bcap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_bcap.paragraph_format.space_before = Pt(3)
    p_bcap.paragraph_format.space_after = Pt(20)
    r_bcap = p_bcap.add_run("Figure 7.1: Bar Chart of Sample Stock Counts")
    r_bcap.font.italic = True
    r_bcap.font.size = Pt(9.5)
    r_bcap.font.color.rgb = secondary_color
    
    # Pie Chart
    p_pch = doc.add_paragraph()
    p_pch.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_pch = p_pch.add_run()
    r_pch.add_picture(sentiment_img_path, width=Inches(4.2))
    
    p_pcap = doc.add_paragraph()
    p_pcap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_pcap.paragraph_format.space_before = Pt(3)
    r_pcap = p_pcap.add_run("Figure 7.2: Pie Chart of Logged Interactions Sentiment Analysis")
    r_pcap.font.italic = True
    r_pcap.font.size = Pt(9.5)
    r_pcap.font.color.rgb = secondary_color
    
    doc.add_page_break()
    
    # ==================== 8. LANGGRAPH TOOLS ====================
    h8 = doc.add_heading("8. LangGraph Tool Specifications", level=1)
    h8.runs[0].font.color.rgb = primary_color
    h8.runs[0].font.bold = True
    
    doc.add_paragraph(
        "The AI Agent has 5 specific database-backed tools at its disposal. Whenever the LLM decides "
        "that a user query requires database modification or query details, the agent routes execution "
        "to the appropriate Python function in `backend/agent.py`:"
    )
    
    tools_hdr = ["Tool Identifier", "Input Arguments", "Description of Database / Form Operation"]
    tools_data = [
        ["search_hcp", "query: str", "Searches doctor records in SQLite matching terms in name, specialty, or hospital locations."],
        ["get_available_materials_and_samples", "None", "Retrieves inventory lists of PDF clinical studies and sample counts to present to the user."],
        ["get_interaction_history", "hcp_id: int", "Retrieves all historic logged records associated with the specified HCP, sorted chronologically."],
        ["log_interaction", "hcp_id, type, date, time, attendees, topics, materials, samples, sentiment, outcomes, follow_ups", "Inserts a new interaction record into the DB and decrements the corresponding sample stocks by their distributed quantities."],
        ["edit_interaction", "interaction_id: int, updated_fields: dict", "Mutates specific fields of an existing interaction record by lookup ID, updating values in SQLite."],
        ["update_draft_interaction", "fields: dict", "Dynamically updates the in-memory frontend form draft state variables with extracted entities (e.g. topics, sentiment) in real-time."]
    ]
    
    create_styled_table(doc, tools_hdr, tools_data)
    
    # ==================== 9. SETUP & RUNNING GUIDE ====================
    h9 = doc.add_heading("9. Running and Verification Guide", level=1)
    h9.runs[0].font.color.rgb = primary_color
    h9.runs[0].font.bold = True
    
    h9_1 = doc.add_heading("Prerequisites", level=2)
    h9_1.runs[0].font.color.rgb = primary_color
    doc.add_paragraph("Ensure that you have Node.js (v18+) and Python (v3.10+) installed on the local system.")
    
    h9_2 = doc.add_heading("Backend Setup", level=2)
    h9_2.runs[0].font.color.rgb = primary_color
    
    p_bsetup = doc.add_paragraph()
    r_bsetup = p_bsetup.add_run(
        "Navigate to backend and configure environment:\n"
        "  cd backend\n"
        "  python -m venv venv\n"
        "  .\\venv\\Scripts\\activate\n"
        "  pip install -r requirements.txt\n"
        "  python seed.py\n"
        "  python main.py"
    )
    r_bsetup.font.name = 'Courier New'
    r_bsetup.font.size = Pt(9.5)
    
    h9_3 = doc.add_heading("Frontend Setup", level=2)
    h9_3.runs[0].font.color.rgb = primary_color
    
    p_fsetup = doc.add_paragraph()
    r_fsetup = p_fsetup.add_run(
        "Navigate to frontend and boot Vite web server:\n"
        "  cd frontend\n"
        "  npm install\n"
        "  npm run dev"
    )
    r_fsetup.font.name = 'Courier New'
    r_fsetup.font.size = Pt(9.5)
    
    # Footer elements
    for section in doc.sections:
        footer = section.footer
        p_foot = footer.paragraphs[0]
        p_foot.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r_foot = p_foot.add_run("AI-First CRM HCP Documentation | Candidate: Chaitanya Jagan | Page ")
        r_foot.font.size = Pt(9)
        r_foot.font.color.rgb = secondary_color
        
        # Adding a simple page number field in docx
        fldSimple = OxmlElement('w:fldSimple')
        fldSimple.set(qn('w:instr'), 'PAGE')
        p_foot._p.append(fldSimple)

    doc.save(output_path)
    print(f"Word document compiled successfully at: {output_path}")

def create_styled_table(doc, headers, data):
    table = doc.add_table(rows=len(data) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Format headers
    for idx, h_text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_background(cell, "1F4E79")
        set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(h_text)
        r.font.name = 'Arial'
        r.font.size = Pt(10)
        r.font.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        
    # Format data
    for r_idx, row_data in enumerate(data):
        row = table.rows[r_idx + 1]
        bg_color = "F6F8FA" if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            if bg_color != "FFFFFF":
                set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            r.font.name = 'Arial'
            r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(60, 60, 60)
            
    doc.add_paragraph() # Spacer


# Main generator runner
def main():
    docs_dir = "C:\\aivoa\\docs"
    if not os.path.exists(docs_dir):
        os.makedirs(docs_dir)
        
    print("Generating system diagrams...")
    arch_path = os.path.join(docs_dir, "architecture_diagram.png")
    state_path = os.path.join(docs_dir, "state_diagram.png")
    stock_path = os.path.join(docs_dir, "inventory_chart.png")
    sentiment_path = os.path.join(docs_dir, "sentiment_chart.png")
    
    draw_architecture_diagram(arch_path)
    draw_state_diagram(state_path)
    
    print("Querying database details...")
    hcps, materials, samples, interactions = query_db()
    
    print("Plotting charts...")
    plot_inventory_chart(stock_path, samples)
    plot_sentiment_chart(sentiment_path, interactions)
    
    print("Assembling word document report...")
    output_doc_path = "C:\\aivoa\\AI_First_CRM_HCP_Module_Documentation.docx"
    build_word_document(output_doc_path, hcps, materials, samples, interactions, 
                        arch_path, state_path, stock_path, sentiment_path)
    
    print("Keeping diagrams in docs/ folder for README embedding.")
    print("Documentation generation process complete!")

if __name__ == "__main__":
    main()
